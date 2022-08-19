import docx 
import pandas as pd

# filename = input('filename: ')
# filename = 'someText.txt'
filename = 'someDoc.docx'
# filename = 'someSheet.xlsx'


class Reader:
    def __init__(self, file):
        self.__file=file

    @property
    def filename(self):
        return self.__file
   
    @filename.setter
    def filename(self,file):
        self.__file=file
    
    @staticmethod
    def file_reader_excel(file):
        table = pd.read_excel(file)
        tdata  = pd.DataFrame(table)
        return f"{tdata}"

    @staticmethod
    def file_reader_doc(file):
        text = []
        t = docx.Document(file)
        for word in t.paragraphs:
            text.append(word.text)
        return f"{text}"

    @staticmethod
    def file_reader_txt(file):
        text = []
        with open(file, 'r') as buffer:
            for line in buffer:
                text.append(line)
        return f"{text}"

class Writer(Reader):
    def __init__(self, file, text):
        Reader.__init__(self,file)
        self.text = text


    @staticmethod
    def file_write_txt(file,text):
        for i in range(len(text)):
            if isinstance(text[i],int) or isinstance(text[i], float):
                text[i] = str(text[i])
        with open(file, 'a') as buffer:
            buffer.write(' '.join(text))
        return f"Source has been updated: \n{Reader.file_reader_txt(file)}"
    
    @staticmethod
    def file_write_docx(file,text):
        doc = docx.Document(file)
        for i in range(len(text)):
            if isinstance(text[i],int) or isinstance(text[i], float):
                text[i] = str(text[i])
        doc.add_paragraph(' '.join(text))
        doc.save(file)
        return f"Source has been updated: \n{Reader.file_reader_doc(file)}"
            
    
    @staticmethod
    def file_write_excel(file,text):
        pass

text = ['asdasd',12,'ssadsad']
writer = Writer(filename,text)
reader = Reader(filename)

# print(reader.file_reader_doc(filename))
# print(reader.file_reader_excel(filename))
# print(writer.file_write_txt(filename,text))
print(writer.file_write_docx(filename,text))
# print(reader.file_reader_txt(filename))
